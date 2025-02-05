import pandas as pd

from docx import Document

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement, ns

# Cargar datos de los documentos Excel
sesiones_df = pd.read_excel(r"C:/Users/famil/Documents/Empresas/Meraki/CLIENTES/REGISTRO DE CONSULTAS.xlsx", sheet_name = "Hoja1")    
clientes_df = pd.read_excel(r"C:/Users/famil/Documents/Empresas/Meraki/CLIENTES/GESTION DE CLIENTES.xlsx", sheet_name = "Hoja1")

# Buscar datos ultima fila registro de sesiones
ultima_sesion = sesiones_df.iloc[-1]

# Buscar nombre cliente, fecha, numero de sesiones y numero de factura de ultima sesion
nombre_cliente = ultima_sesion['NOMBRES Y APELLIDOS']
fecha_factura = ultima_sesion['FECHA'].strftime('%d-%m-%Y')
numero_sesiones = ultima_sesion["N SESIONES"]
numero_factura = ultima_sesion['FACTURA']

# Calcular los importes ultima sesion (IVA exento)
base_imponible = 50 * numero_sesiones
cuota_integra  = 0.21 * base_imponible
total = base_imponible + cuota_integra

# Buscar datos en registro de clientes
cliente_info = clientes_df.loc[clientes_df['NOMBRES Y APELLIDOS'] == nombre_cliente]

# Obtener la informacion del cliente (direccion, DNI y telefono)
if not cliente_info.empty:
    direccion_cliente = cliente_info['DIRECCION'].iloc[0]
    dni_cliente = cliente_info['IDENTIFICACION'].iloc[0]
    telefono_cliente = int(cliente_info['TELEFONO'].iloc[0])
else:
    direccion_cliente = ''
    dni_cliente = ''
    telefono_cliente = ''

# Informacion de la clinica
nombre_colegiada = "name and surname"
dni_colegiada = "ID"
numero_colegiada = "collegiate number"
direccion_colegiada = "address"
telefono_colegiada = "phone number"
correo_colegiada = "email"

# Factura como diccionario
factura = {
    "encabezado": "FACTURA",
    "fecha": fecha_factura,
    "n_factura": numero_factura,
    "nombre_1": nombre_colegiada,
    "numero_de_identificacion_1": dni_colegiada,
    "n_colegiada": numero_colegiada,
    "direccion_1": direccion_colegiada,
    "telefono_1": telefono_colegiada,
    "correo_1": correo_colegiada,
    "nombre_2": nombre_cliente,
    "direccion_2": direccion_cliente,
    "numero_de_identificacion_2": dni_cliente,
    "telefono_2": telefono_cliente,
    "descripcion": "ASISTENCIA PSICOLOGICA",
    "n_sesiones": numero_sesiones,
    "base_imponible": base_imponible,
    "cuota_integra_(21%)": cuota_integra,
    "total": total,
    "exencion": "* Exenta de IVA, según Ley 37/1992, artículo 20, número 3º.", 
}

# Generar factura como documento Word
doc = Document()

estilo_fuente = 'Franklin Gothic Demi'
tamano_fuente = 12

    # Encabezado
parrafo_encabezado = doc.add_paragraph('FACTURA')
parrafo_encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = parrafo_encabezado.runs[0]
run.font.size = Pt(14)
run.font.name = estilo_fuente
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)

    # Fecha y Numero de factura
parrafo_fecha_factura = doc.add_paragraph()
run_fecha_factura = parrafo_fecha_factura.add_run(f'FECHA: {factura["fecha"]}')
run_fecha_factura.font.size = Pt(tamano_fuente)
run_fecha_factura.font.name = estilo_fuente
run_fecha_factura.bold = True
parrafo_fecha_factura.alignment = WD_ALIGN_PARAGRAPH.LEFT

parrafo_numero_factura = doc.add_paragraph()
run_numero_factura = parrafo_numero_factura.add_run(f'N FACTURA: {factura["n_factura"]}')
run_numero_factura.font.size = Pt(tamano_fuente)
run_numero_factura.font.name = estilo_fuente
run_numero_factura.bold = True
parrafo_numero_factura.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Imagen
doc.add_picture(r"C:/Users/famil/Documents/Empresas/Meraki/CLIENTES/LOGO.png", height= Inches(1.2), width= Inches(1.2))
parrafo_imagen = doc.paragraphs[-1]
parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Informacion Despacho
info_despacho = [
    factura['nombre_1'],
    factura['direccion_1'],
    f'DNI: {factura["numero_de_identificacion_1"]}',
    f'TELEFONO: {factura["telefono_1"]}',
    factura['n_colegiada'],
    factura['correo_1']
]

for info in info_despacho:
    parrafo_despacho = doc.add_paragraph()
    run_despacho = parrafo_despacho.add_run(info)
    parrafo_despacho.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_despacho.font.size = Pt(tamano_fuente)
    run_despacho.font.name = estilo_fuente
    run_despacho.bold = True

    # Informacion Clientes
info_cliente = [
    factura['nombre_2'],
    factura['direccion_2'],
    f'DNI: {factura["numero_de_identificacion_2"]}',
    f'TELEFONO: {factura["telefono_2"]}'
]

for info in info_cliente:
    parrafo_cliente = doc.add_paragraph()
    run_cliente = parrafo_cliente.add_run(info)
    parrafo_cliente.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_cliente.font.size = Pt(tamano_fuente)
    run_cliente.font.name = estilo_fuente
    run_cliente.bold = True

    # Tabla
tabla = doc.add_table(rows=3, cols=5)
tabla.cell(0, 0).text = "DESCRIPCION"
tabla.cell(0, 1).text = "P SESION"
tabla.cell(0, 2).text = "N SESIONES"
tabla.cell(0, 3).text = "I.V.A. (21%)"
tabla.cell(0, 4).text = "TOTAL"
tabla.cell(1, 0).text = factura["descripcion"]
tabla.cell(1, 1).text = "50€"
tabla.cell(1, 2).text = str(factura["n_sesiones"])
tabla.cell(1, 3).text = str(factura["cuota_integra_(21%)"]) + "€ *"
tabla.cell(1, 4).text = str(factura["total"]) + "€"
tabla.cell(2, 3).text = "TOTAL FACTURA"
tabla.cell(2, 4).text = str(factura["base_imponible"]) + "€"

for row in tabla.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # Buscar el elemento 'tcBorders' usando el espacio de nombres
        tcBorders = tcPr.find(ns.qn('w:tcBorders'))
        
        # Si no existe, crearlo
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        
        # Configurar bordes para cada lado de la celda
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = tcBorders.find(ns.qn(f'w:{border_name}'))
            if border is None:
                border = OxmlElement(f'w:{border_name}')
                tcBorders.append(border)
            # Configurar propiedades del borde
            border.set(ns.qn('w:val'), 'single')
            border.set(ns.qn('w:sz'), '4')
            border.set(ns.qn('w:space'), '0')
            border.set(ns.qn('w:color'), '000000')
        
        # Alinear texto y establecer estilo
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(tamano_fuente)
                run.font.name = estilo_fuente
                run.bold = True

    # Exencion
parrafo_exencion = doc.add_paragraph()
run_parrafo_exencion = parrafo_exencion.add_run(factura['exencion'])
run_parrafo_exencion.font.size = Pt(tamano_fuente)
run_parrafo_exencion.font.name = estilo_fuente
run_parrafo_exencion.bold = True
parrafo_exencion.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Guardar documento
ruta_factura = r"C:/Users/famil/Documents/Empresas/Meraki/CLIENTES/FACTURA X.docx"
doc.save(ruta_factura)